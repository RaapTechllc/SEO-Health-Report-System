#!/usr/bin/env python3
"""
SEO Health Report Generator for Sheet Metal Werks
Runs all three audits and generates a comprehensive report.
"""

import os
import sys
from dotenv import load_dotenv

# Load environment variables first
load_dotenv(".env.local")

# Add project root to path
project_root = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, project_root)

# Add each module folder to path so relative imports work
for folder in [
    "seo-technical-audit",
    "seo-content-authority",
    "ai-visibility-audit",
    "seo-health-report",
]:
    folder_path = os.path.join(project_root, folder)
    if folder_path not in sys.path:
        sys.path.insert(0, folder_path)

# Create module aliases for hyphenated folder names
import importlib.util


def register_module_alias(alias_name: str, folder_name: str):
    """Register a module alias for a hyphenated folder."""
    folder_path = os.path.join(project_root, folder_name)
    init_path = os.path.join(folder_path, "__init__.py")
    if os.path.exists(init_path):
        spec = importlib.util.spec_from_file_location(
            alias_name, init_path, submodule_search_locations=[folder_path]
        )
        module = importlib.util.module_from_spec(spec)
        sys.modules[alias_name] = module
        try:
            spec.loader.exec_module(module)
        except Exception:
            pass  # Some modules may have import errors, but we still register them


def register_scripts_submodule(parent_alias: str, folder_name: str):
    """Register the scripts submodule for a hyphenated folder."""
    folder_path = os.path.join(project_root, folder_name, "scripts")
    init_path = os.path.join(folder_path, "__init__.py")
    submodule_name = f"{parent_alias}.scripts"
    if os.path.exists(init_path):
        spec = importlib.util.spec_from_file_location(
            submodule_name, init_path, submodule_search_locations=[folder_path]
        )
        module = importlib.util.module_from_spec(spec)
        sys.modules[submodule_name] = module
        try:
            spec.loader.exec_module(module)
        except Exception:
            pass


# Register the main modules and their scripts submodules
# Order matters: seo_health_report first as it has dependencies
register_module_alias("seo_health_report", "seo-health-report")
register_scripts_submodule("seo_health_report", "seo-health-report")

register_module_alias("seo_technical_audit", "seo-technical-audit")
register_scripts_submodule("seo_technical_audit", "seo-technical-audit")

register_module_alias("seo_content_authority", "seo-content-authority")
register_scripts_submodule("seo_content_authority", "seo-content-authority")

register_module_alias("ai_visibility_audit", "ai-visibility-audit")
register_scripts_submodule("ai_visibility_audit", "ai-visibility-audit")

# Now we can import from the scripts folders directly
from datetime import datetime
from typing import Dict, Any, List

# Configuration for Sheet Metal Werks
TARGET_URL = "https://www.sheetmetalwerks.com/"
COMPANY_NAME = "Sheet Metal Werks"
PRIMARY_KEYWORDS = [
    "sheet metal fabrication",
    "custom metal fabrication",
    "metal stamping",
    "precision sheet metal",
    "sheet metal manufacturing",
    "metal forming services",
    "sheet metal prototyping",
    "industrial metal fabrication",
]
GROUND_TRUTH = {
    "company_name": "Sheet Metal Werks",
    "industry": "Sheet Metal Fabrication / Manufacturing",
    "services": [
        "Sheet metal fabrication",
        "Custom metal fabrication",
        "Metal stamping",
    ],
    "location": "United States",
}

# Competitor Analysis (Top 3-5 competitors for benchmarking)
COMPETITORS = [
    "metalfabricating.com",
    "precisionmetalworks.com",
    "custommetalfab.com",
    "industrialsheetmetal.com",
]

# Business Metrics (for ROI calculations)
BUSINESS_METRICS = {
    "monthly_revenue": 150000,  # Current monthly revenue
    "avg_order_value": 2500,  # Average project value
    "conversion_rate": 0.03,  # 3% website conversion rate
    "traffic_increase": "25",  # Estimated traffic increase %
    "revenue_opportunity": "18,000",  # Monthly revenue opportunity
    "competitive_advantage": "AI visibility + technical expertise positioning",
}

# Logo configuration
CLIENT_LOGO = os.path.join(project_root, "reports", "logos", "sheetmetalwerks_logo.png")
AGENCY_LOGO = os.path.join(project_root, "reports", "logos", "raaptech_logo.png")
AGENCY_NAME = "Raaptech"
BRAND_COLORS = {"primary": "#1a365d", "secondary": "#2b6cb0"}


def run_technical_audit() -> Dict[str, Any]:
    """Run the technical SEO audit."""
    print("\n[1/3] Running Technical SEO Audit...")
    try:
        import asyncio

        # Import from properly registered modules
        from seo_technical_audit.scripts.crawl_site import analyze_crawlability
        from seo_technical_audit.scripts.analyze_speed import analyze_speed
        from seo_technical_audit.scripts.check_security import analyze_security
        from seo_technical_audit.scripts.validate_schema import validate_structured_data

        results = {"score": 0, "grade": "F", "components": {}, "recommendations": []}

        # Crawlability (20 points)
        crawl = analyze_crawlability(TARGET_URL, depth=50)
        results["components"]["crawlability"] = crawl

        # Indexing (15 points)
        results["components"]["indexing"] = {
            "score": min(15, crawl.get("sitemaps", {}).get("score", 7)),
            "max": 15,
            "findings": [],
        }

        # Speed (25 points) - handle async function
        speed_coro = analyze_speed(TARGET_URL, strategy="mobile")
        speed = asyncio.run(speed_coro)
        results["components"]["speed"] = speed

        # Mobile (15 points)
        mobile_psi = speed.get("psi_score") if isinstance(speed, dict) else None
        mobile_psi = mobile_psi if mobile_psi is not None else 50
        mobile_score = (
            15
            if mobile_psi >= 90
            else 12
            if mobile_psi >= 75
            else 9
            if mobile_psi >= 50
            else 6
        )
        results["components"]["mobile"] = {
            "score": mobile_score,
            "max": 15,
            "psi_score": mobile_psi,
        }

        # Security (10 points)
        security = analyze_security(TARGET_URL)
        results["components"]["security"] = security

        # Structured Data (15 points)
        schema = validate_structured_data(TARGET_URL)
        results["components"]["structured_data"] = schema

        # Calculate total
        results["score"] = sum(
            c.get("score", 0) for c in results["components"].values()
        )
        results["grade"] = (
            "A"
            if results["score"] >= 90
            else "B"
            if results["score"] >= 80
            else "C"
            if results["score"] >= 70
            else "D"
            if results["score"] >= 60
            else "F"
        )

        print(f"    Technical Score: {results['score']}/100 ({results['grade']})")
        return results
    except Exception as e:
        print(f"    Technical audit error: {e}")
        import traceback

        traceback.print_exc()
        return {"score": None, "grade": "N/A", "status": "unavailable", "message": "Technical audit unavailable - contact RaapTech", "components": {}, "error": str(e)}


def run_content_audit() -> Dict[str, Any]:
    """Run the content & authority audit."""
    print("\n[2/3] Running Content & Authority Audit...")
    try:
        import re
        from urllib.parse import urlparse, urljoin

        # Import from properly registered modules
        from seo_content_authority.scripts.analyze_content import (
            fetch_page,
            analyze_page_content,
            assess_content_quality,
        )
        from seo_content_authority.scripts.check_eeat import analyze_eeat_signals
        from seo_content_authority.scripts.map_topics import analyze_topical_coverage
        from seo_content_authority.scripts.analyze_links import analyze_internal_links
        from seo_content_authority.scripts.score_backlinks import (
            analyze_backlink_profile,
            estimate_backlink_health,
        )

        results = {
            "score": 0,
            "grade": "F",
            "components": {},
            "content_gaps": [],
            "recommendations": [],
        }

        # Crawl pages
        html = fetch_page(TARGET_URL)
        pages = [TARGET_URL]
        if html:
            links = re.findall(r'href=["\']([^"\']+)["\']', html, re.IGNORECASE)
            parsed_base = urlparse(TARGET_URL)
            for link in links[:30]:
                full_url = urljoin(TARGET_URL, link)
                parsed = urlparse(full_url)
                if parsed.netloc == parsed_base.netloc and full_url not in pages:
                    pages.append(full_url)

        # Content Quality (25 points)
        page_analyses = [analyze_page_content(url) for url in pages[:20]]
        content = assess_content_quality(page_analyses)
        results["components"]["content_quality"] = content

        # E-E-A-T (20 points)
        eeat = analyze_eeat_signals(TARGET_URL)
        results["components"]["eeat"] = eeat

        # Keyword Position (15 points) - neutral without ranking API
        results["components"]["keyword_position"] = {
            "score": 7,
            "max": 15,
            "note": "Requires ranking API",
        }

        # Topical Authority (15 points)
        topics = analyze_topical_coverage(TARGET_URL, PRIMARY_KEYWORDS, 30)
        results["components"]["topical_authority"] = topics
        results["content_gaps"] = topics.get("content_gaps", [])

        # Backlinks (15 points)
        backlinks = analyze_backlink_profile(TARGET_URL)
        if backlinks.get("score", 0) == 0:
            backlinks = estimate_backlink_health(TARGET_URL)
        results["components"]["backlinks"] = backlinks

        # Internal Links (10 points)
        internal = analyze_internal_links(TARGET_URL, 30)
        results["components"]["internal_links"] = internal

        # Calculate total
        results["score"] = sum(
            c.get("score", 0) for c in results["components"].values()
        )
        results["grade"] = (
            "A"
            if results["score"] >= 90
            else "B"
            if results["score"] >= 80
            else "C"
            if results["score"] >= 70
            else "D"
            if results["score"] >= 60
            else "F"
        )

        print(f"    Content Score: {results['score']}/100 ({results['grade']})")
        return results
    except Exception as e:
        print(f"    Content audit error: {e}")
        import traceback

        traceback.print_exc()
        return {"score": None, "grade": "N/A", "status": "unavailable", "message": "Content audit unavailable - contact RaapTech", "components": {}, "error": str(e)}


def run_ai_visibility_audit() -> Dict[str, Any]:
    """Run the AI visibility audit."""
    print("\n[3/3] Running AI Visibility Audit...")
    try:
        import asyncio

        # Import from properly registered modules
        from ai_visibility_audit.scripts.query_ai_systems import (
            generate_test_queries,
            query_all_systems,
        )
        from ai_visibility_audit.scripts.analyze_responses import (
            analyze_brand_presence,
            check_accuracy,
            analyze_sentiment,
        )
        from ai_visibility_audit.scripts.check_parseability import (
            analyze_site_structure,
        )
        from ai_visibility_audit.scripts.check_knowledge import check_all_sources
        from ai_visibility_audit.scripts.score_citability import (
            analyze_content_citability,
        )

        results = {
            "score": 0,
            "grade": "F",
            "components": {},
            "ai_responses": [],
            "recommendations": [],
        }

        # Generate and run queries (handle async)
        queries = generate_test_queries(COMPANY_NAME, PRIMARY_KEYWORDS, [], None)

        # query_all_systems is async
        responses_coro = query_all_systems(
            queries, COMPANY_NAME, ["claude", "openai", "perplexity"]
        )
        responses = asyncio.run(responses_coro)

        # Store responses
        for system, sys_responses in responses.items():
            for r in sys_responses:
                if not r.error:
                    results["ai_responses"].append(
                        {
                            "query": r.query,
                            "system": r.system,
                            "response": r.response,
                            "brand_mentioned": r.brand_mentioned,
                            "position": r.position,
                        }
                    )

        # AI Presence (25 points)
        presence = analyze_brand_presence(responses, COMPANY_NAME, [])
        results["components"]["ai_presence"] = presence

        # Accuracy (20 points)
        accuracy = (
            check_accuracy(responses, GROUND_TRUTH, COMPANY_NAME)
            if GROUND_TRUTH
            else {"score": 10, "max": 20}
        )
        results["components"]["accuracy"] = accuracy

        # Parseability (15 points)
        parseability = analyze_site_structure(TARGET_URL)
        results["components"]["parseability"] = parseability

        # Knowledge Graph (15 points)
        knowledge = check_all_sources(COMPANY_NAME, TARGET_URL)
        results["components"]["knowledge_graph"] = knowledge

        # Citation Likelihood (15 points)
        citation = analyze_content_citability(TARGET_URL)
        results["components"]["citation_likelihood"] = citation

        # Sentiment (10 points)
        sentiment = analyze_sentiment(responses, COMPANY_NAME)
        results["components"]["sentiment"] = sentiment

        # Calculate total
        results["score"] = sum(
            c.get("score", 0) for c in results["components"].values()
        )
        results["grade"] = (
            "A"
            if results["score"] >= 90
            else "B"
            if results["score"] >= 80
            else "C"
            if results["score"] >= 70
            else "D"
            if results["score"] >= 60
            else "F"
        )

        print(f"    AI Visibility Score: {results['score']}/100 ({results['grade']})")
        return results
    except Exception as e:
        print(f"    AI visibility audit error: {e}")
        import traceback

        traceback.print_exc()
        return {"score": None, "grade": "N/A", "status": "unavailable", "message": "AI audit unavailable - contact RaapTech", "components": {}, "error": str(e)}


def calculate_overall_score(technical: Dict, content: Dict, ai: Dict) -> Dict[str, Any]:
    """Calculate weighted overall score."""
    tech_score = technical.get("score")
    content_score = content.get("score")
    ai_score = ai.get("score")

    # Weights: Technical 30%, Content 35%, AI 35%
    overall = (tech_score * 0.30) + (content_score * 0.35) + (ai_score * 0.35)
    overall = round(overall)

    grade = (
        "A"
        if overall >= 90
        else "B"
        if overall >= 80
        else "C"
        if overall >= 70
        else "D"
        if overall >= 60
        else "F"
    )

    return {
        "overall_score": overall,
        "grade": grade,
        "component_scores": {
            "technical": {"score": tech_score, "weight": 0.30},
            "content": {"score": content_score, "weight": 0.35},
            "ai_visibility": {"score": ai_score, "weight": 0.35},
        },
    }


def print_report(overall: Dict, technical: Dict, content: Dict, ai: Dict):
    """Print a formatted summary report."""
    print("\n" + "=" * 60)
    print("SEO HEALTH REPORT SUMMARY")
    print("=" * 60)
    print(f"\nCompany: {COMPANY_NAME}")
    print(f"URL: {TARGET_URL}")
    print(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    print(f"\n{'=' * 60}")
    print(f"OVERALL SCORE: {overall['overall_score']}/100 (Grade: {overall['grade']})")
    print("=" * 60)

    print("\nCOMPONENT SCORES:")
    print("-" * 40)
    print(f"  Technical SEO:     {technical.get('score', 'N/A')}/100 (30% weight)")
    print(f"  Content Authority: {content.get('score', 'N/A')}/100 (35% weight)")
    print(f"  AI Visibility:     {ai.get('score', 'N/A')}/100 (35% weight)")

    # Print key findings from each audit
    print("\n" + "-" * 60)
    print("KEY FINDINGS")
    print("-" * 60)

    for name, audit in [
        ("Technical", technical),
        ("Content", content),
        ("AI Visibility", ai),
    ]:
        components = audit.get("components", {})
        if components:
            print(f"\n{name}:")
            for comp_name, comp_data in list(components.items())[:3]:
                if isinstance(comp_data, dict):
                    score = comp_data.get("score", "?")
                    max_score = comp_data.get("max", 100)
                    print(
                        f"  - {comp_name.replace('_', ' ').title()}: {score}/{max_score}"
                    )

    print("\n" + "=" * 60)


def main():
    """Main entry point."""
    print("=" * 60)
    print("SEO HEALTH REPORT GENERATOR")
    print("=" * 60)
    print(f"Target: {TARGET_URL}")
    print(f"Company: {COMPANY_NAME}")
    print(f"Keywords: {', '.join(PRIMARY_KEYWORDS[:5])}...")
    print(f"Prepared by: {AGENCY_NAME}")
    print("=" * 60)

    # Run all audits from project root (no directory changes needed now)
    technical = run_technical_audit()
    content = run_content_audit()
    ai = run_ai_visibility_audit()

    # Calculate overall score
    overall = calculate_overall_score(technical, content, ai)


    # Add ROI analysis
    try:
        import sys
        sys.path.insert(0, os.path.join(project_root, "seo-health-report", "scripts"))
        from roi_calculator import calculate_seo_roi, generate_roi_summary
        
        roi_projections = calculate_seo_roi(
            {"audits": {"technical": technical, "content": content, "ai_visibility": ai}},
            BUSINESS_METRICS
        )
        roi_summary = generate_roi_summary(roi_projections)
        
        overall["roi_analysis"] = roi_summary
        
        print(f"\n💰 ROI ANALYSIS:")
        print(f"    Total Investment: ${roi_summary["total_investment"]:,}")
        print(f"    Annual Revenue Potential: ${roi_summary["total_annual_revenue"]:,}")
        print(f"    Overall ROI: {roi_summary["overall_roi"]}%")
        print(f"    Payback Period: {roi_summary["payback_months"]} months")
        
    except Exception as e:
        print(f"    ROI calculation error: {e}")
        overall["roi_analysis"] = None
    # Print report
    print_report(overall, technical, content, ai)

    # Save results to JSON
    import json

    results = {
        "company": COMPANY_NAME,
        "url": TARGET_URL,
        "timestamp": datetime.now().isoformat(),
        "prepared_by": AGENCY_NAME,
        "overall": overall,
        "technical": technical,
        "content": content,
        "ai_visibility": ai,
    }

    output_file = os.path.join(
        project_root,
        "reports",
        f"sheetmetalwerks-audit-{datetime.now().strftime('%Y%m%d')}.json",
    )
    with open(output_file, "w") as f:
        json.dump(results, f, indent=2, default=str)
    print(f"\nDetailed results saved to: {output_file}")

    # Generate DOCX report with logos
    generate_docx_report(results, overall, technical, content, ai)

    return results


def generate_docx_report(
    results: Dict, overall: Dict, technical: Dict, content: Dict, ai: Dict
):
    """Generate a premium branded DOCX report with AI focus."""
    try:
        # Import the premium template
        import sys
        import os

        # Add the scripts directory to path
        scripts_dir = os.path.join(project_root, "seo-health-report", "scripts")
        if scripts_dir not in sys.path:
            sys.path.insert(0, scripts_dir)

        from premium_report_template import generate_premium_docx_report

        # Use premium template with proper branding
        output_path = generate_premium_docx_report(
            results=results,
            overall=overall,
            technical=technical,
            content=content,
            ai=ai,
            company_name=COMPANY_NAME,
            target_url=TARGET_URL,
            agency_name=AGENCY_NAME,
            client_logo_path=CLIENT_LOGO,
            agency_logo_path=AGENCY_LOGO,
        )

        if output_path:
            print(f"\n[PREMIUM REPORT] {output_path}")
            print("Professional branded report ready for client delivery")
            return

    except Exception as e:
        print(f"\n[WARNING] Premium template failed, using basic template: {e}")

    # Fallback to basic template - Generate a branded DOCX report with logos.
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Cover Page with logos
        # Add agency logo (top right)
        if os.path.exists(AGENCY_LOGO):
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header_para.add_run()
            run.add_picture(AGENCY_LOGO, width=Inches(1.5))

        # Add client logo (centered)
        if os.path.exists(CLIENT_LOGO):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(CLIENT_LOGO, width=Inches(3))

        # Title
        title = doc.add_heading("SEO Health Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{COMPANY_NAME}")
        run.bold = True
        run.font.size = Pt(18)

        # Date and preparer
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"\nPrepared by: {AGENCY_NAME}\n")
        info.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")

        doc.add_page_break()

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)

        grade_desc = {
            "A": "Excellent - Your SEO health is outstanding!",
            "B": "Good - Strong foundation with room for improvement",
            "C": "Needs Work - Several areas require attention",
            "D": "Poor - Significant improvements needed",
            "F": "Critical - Urgent action required",
        }

        summary = doc.add_paragraph()
        summary.add_run(
            f"Overall Score: {overall['overall_score']}/100 (Grade: {overall['grade']})\n\n"
        ).bold = True
        summary.add_run(
            f"{grade_desc.get(overall['grade'], 'Assessment complete')}\n\n"
        )

        # Score breakdown table
        doc.add_heading("Score Breakdown", level=2)
        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Component"
        headers[1].text = "Score"
        headers[2].text = "Weight"

        rows_data = [
            ("Technical SEO", f"{technical.get('score', 'N/A')}/100", "30%"),
            ("Content & Authority", f"{content.get('score', 'N/A')}/100", "35%"),
            ("AI Visibility", f"{ai.get('score', 'N/A')}/100", "35%"),
        ]

        for i, (comp, score, weight) in enumerate(rows_data, 1):
            row = table.rows[i].cells
            row[0].text = comp
            row[1].text = score
            row[2].text = weight

        doc.add_page_break()

        # Technical SEO Section
        doc.add_heading("Technical SEO Analysis", level=1)
        doc.add_paragraph(
            f"Score: {technical.get('score', 'N/A')}/100 ({technical.get('grade', '?')})"
        )

        if technical.get("components"):
            for comp_name, comp_data in technical["components"].items():
                if isinstance(comp_data, dict):
                    doc.add_heading(comp_name.replace("_", " ").title(), level=2)
                    doc.add_paragraph(
                        f"Score: {comp_data.get('score', '?')}/{comp_data.get('max', 100)}"
                    )

        doc.add_page_break()

        # Content & Authority Section
        doc.add_heading("Content & Authority Analysis", level=1)
        doc.add_paragraph(
            f"Score: {content.get('score', 'N/A')}/100 ({content.get('grade', '?')})"
        )

        if content.get("components"):
            for comp_name, comp_data in content["components"].items():
                if isinstance(comp_data, dict):
                    doc.add_heading(comp_name.replace("_", " ").title(), level=2)
                    doc.add_paragraph(
                        f"Score: {comp_data.get('score', '?')}/{comp_data.get('max', 100)}"
                    )

        doc.add_page_break()

        # AI Visibility Section
        doc.add_heading("AI Visibility Analysis", level=1)
        doc.add_paragraph(
            f"Score: {ai.get('score', 'N/A')}/100 ({ai.get('grade', '?')})"
        )
        doc.add_paragraph(
            "\nThis section evaluates how your brand appears in AI-generated responses from systems like ChatGPT, Claude, and Perplexity."
        )

        if ai.get("components"):
            for comp_name, comp_data in ai["components"].items():
                if isinstance(comp_data, dict):
                    doc.add_heading(comp_name.replace("_", " ").title(), level=2)
                    doc.add_paragraph(
                        f"Score: {comp_data.get('score', '?')}/{comp_data.get('max', 100)}"
                    )

        # Footer with agency branding
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"© {datetime.now().year} {AGENCY_NAME} | SEO Health Report for {COMPANY_NAME}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save document
        output_path = os.path.join(
            project_root,
            "reports",
            f"{COMPANY_NAME.replace(' ', '-')}-SEO-Report-{datetime.now().strftime('%Y-%m-%d')}.docx",
        )
        doc.save(output_path)
        print(f"\nDOCX Report saved to: {output_path}")

    except ImportError:
        print(
            "\nNote: python-docx not installed. Run 'pip install python-docx' for DOCX report generation."
        )
    except Exception as e:
        print(f"\nError generating DOCX report: {e}")
        import traceback

        traceback.print_exc()


if __name__ == "__main__":
    main()
