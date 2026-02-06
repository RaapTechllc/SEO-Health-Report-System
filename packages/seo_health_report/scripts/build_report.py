"""
Build Report Document

Assemble complete report document from audit data.
"""

import os
from datetime import datetime
from typing import Any, Optional

from .logger import get_logger

logger = get_logger(__name__)


def build_report_document(
    audit_results: dict[str, Any],
    scores: dict[str, Any],
    executive_summary: dict[str, Any],
    company_name: str,
    logo_file: Optional[str] = None,
    brand_colors: Optional[dict[str, str]] = None,
    output_format: str = "docx",
    output_dir: Optional[str] = None,
    preparer_name: str = "SEO Health Report System",
) -> dict[str, Any]:
    """
    Build the complete report document.

    Args:
        audit_results: Results from run_full_audit()
        scores: Composite scores from calculate_composite_score()
        executive_summary: Summary from generate_executive_summary()
        company_name: Client company name
        logo_file: Path to logo file
        brand_colors: Brand color configuration
        output_format: "docx" or "pdf"
        output_dir: Output directory
        preparer_name: Name to show as report preparer

    Returns:
        Dict with output path and metadata
    """
    result = {
        "success": False,
        "output_path": None,
        "format": output_format,
        "pages": 0,
        "sections": [],
        "error": None,
    }

    # Default output directory
    if not output_dir:
        output_dir = os.getcwd()

    # Generate filename
    date_str = datetime.now().strftime("%Y-%m-%d")
    safe_company = company_name.replace(" ", "-").replace("/", "-")
    filename = f"{safe_company}-SEO-Health-Report-{date_str}.{output_format}"
    output_path = os.path.join(output_dir, filename)

    # Build sections
    sections = []

    # Section 1: Cover Page
    sections.append(
        {
            "type": "cover",
            "content": {
                "title": "SEO Health Report",
                "company_name": company_name,
                "date": datetime.now().strftime("%B %d, %Y"),
                "prepared_by": preparer_name,
                "logo": logo_file,
            },
        }
    )

    # Section 2: Executive Summary
    sections.append({"type": "executive_summary", "content": executive_summary})

    # Section 3: Technical Health
    technical_data = audit_results.get("audits", {}).get("technical", {})
    sections.append(
        {
            "type": "technical",
            "title": "Technical Health",
            "score": scores.get("component_scores", {})
            .get("technical", {})
            .get("score", 0),
            "content": build_technical_section(technical_data),
        }
    )

    # Section 4: Content & Authority
    content_data = audit_results.get("audits", {}).get("content", {})
    sections.append(
        {
            "type": "content",
            "title": "Content & Authority",
            "score": scores.get("component_scores", {})
            .get("content", {})
            .get("score", 0),
            "content": build_content_section(content_data),
        }
    )

    # Section 5: AI Visibility
    ai_data = audit_results.get("audits", {}).get("ai_visibility", {})
    sections.append(
        {
            "type": "ai_visibility",
            "title": "AI Visibility",
            "score": scores.get("component_scores", {})
            .get("ai_visibility", {})
            .get("score", 0),
            "content": build_ai_visibility_section(ai_data),
        }
    )

    # Section 6: Action Plan
    sections.append(
        {
            "type": "action_plan",
            "title": "Action Plan",
            "content": build_action_plan(audit_results, scores),
        }
    )

    # Section 7: Appendix
    sections.append(
        {
            "type": "appendix",
            "title": "Appendix",
            "content": build_appendix(audit_results),
        }
    )

    result["sections"] = sections

    # Generate document based on format
    actual_path = output_path  # Track actual output path (may change on fallback)
    if output_format == "docx":
        success, actual_path = generate_docx(sections, output_path, brand_colors)
    elif output_format == "pdf":
        success = generate_pdf(sections, output_path, brand_colors)
    else:
        # Default to markdown
        success = generate_markdown(sections, output_path)

    if success:
        result["success"] = True
        result["output_path"] = actual_path
        result["format"] = actual_path.split(".")[-1] if actual_path else output_format
        result["pages"] = estimate_page_count(sections)
    else:
        result["error"] = "Document generation failed"

    return result


def build_technical_section(data: dict[str, Any]) -> dict[str, Any]:
    """Build technical health section content."""
    section = {"overview": [], "components": [], "issues": [], "recommendations": []}

    if not data:
        section["overview"].append("Technical audit data not available.")
        return section

    # Overview findings
    section["overview"] = data.get("findings", [])[:5]

    # Component breakdown
    for comp_name, comp_data in data.get("components", {}).items():
        if isinstance(comp_data, dict):
            section["components"].append(
                {
                    "name": comp_name.replace("_", " ").title(),
                    "score": comp_data.get("score", 0),
                    "max": comp_data.get("max", 0),
                    "findings": comp_data.get("findings", [])[:3],
                    "issues": comp_data.get("issues", [])[:3],
                }
            )

    # Top issues
    section["issues"] = data.get("critical_issues", [])[:5]

    # Recommendations
    section["recommendations"] = data.get("recommendations", [])[:5]

    return section


def build_content_section(data: dict[str, Any]) -> dict[str, Any]:
    """Build content & authority section content."""
    section = {
        "overview": [],
        "eeat_assessment": {},
        "content_gaps": [],
        "recommendations": [],
    }

    if not data:
        section["overview"].append("Content audit data not available.")
        return section

    # Overview
    section["overview"] = data.get("findings", [])[:5]

    # E-E-A-T Assessment
    eeat = data.get("components", {}).get("eeat", {})
    section["eeat_assessment"] = {
        "score": eeat.get("score", 0),
        "has_authors": eeat.get("has_authors", False),
        "has_about_page": eeat.get("has_about_page", False),
        "findings": eeat.get("findings", []),
    }

    # Content gaps
    section["content_gaps"] = data.get("content_gaps", [])[:5]

    # Recommendations
    section["recommendations"] = data.get("recommendations", [])[:5]

    return section


def build_ai_visibility_section(data: dict[str, Any]) -> dict[str, Any]:
    """Build AI visibility section content."""
    section = {
        "overview": [],
        "ai_responses": [],
        "accuracy_issues": [],
        "knowledge_graph": {},
        "recommendations": [],
    }

    if not data:
        section["overview"].append("AI visibility audit data not available.")
        return section

    # Overview
    section["overview"].append(f"AI Visibility Score: {data.get('score', 0)}/100")

    # Sample AI responses
    for response in data.get("ai_responses", [])[:3]:
        section["ai_responses"].append(
            {
                "query": response.get("query", ""),
                "system": response.get("system", ""),
                "brand_mentioned": response.get("brand_mentioned", False),
                "sentiment": response.get("sentiment", "neutral"),
                "excerpt": response.get("response", "")[:200] + "...",
            }
        )

    # Accuracy issues
    section["accuracy_issues"] = data.get("accuracy_issues", [])[:5]

    # Knowledge graph status
    kg = data.get("components", {}).get("knowledge_graph", {})
    section["knowledge_graph"] = {
        "score": kg.get("score", 0),
        "sources": kg.get("sources", {}),
        "findings": kg.get("findings", []),
    }

    # Recommendations
    section["recommendations"] = data.get("recommendations", [])[:5]

    return section


def build_action_plan(
    audit_results: dict[str, Any], scores: dict[str, Any]
) -> dict[str, Any]:
    """Build action plan section content."""
    plan = {"quick_wins": [], "strategic_initiatives": [], "prioritized_tasks": []}

    # Collect all recommendations
    all_recs = []
    for audit_name, audit_data in audit_results.get("audits", {}).items():
        if audit_data:
            for rec in audit_data.get("recommendations", []):
                rec_copy = rec.copy()
                rec_copy["source"] = audit_name
                all_recs.append(rec_copy)

    # Categorize
    for rec in all_recs:
        impact = rec.get("impact", "medium")
        effort = rec.get("effort", "medium")

        if impact in ["high", "medium"] and effort == "low":
            plan["quick_wins"].append(rec)
        elif impact == "high" and effort in ["medium", "high"]:
            plan["strategic_initiatives"].append(rec)

    # Sort by priority
    priority_order = {"high": 0, "medium": 1, "low": 2}
    all_recs.sort(key=lambda x: priority_order.get(x.get("priority", "low"), 99))
    plan["prioritized_tasks"] = all_recs[:20]

    # Limit lists
    plan["quick_wins"] = plan["quick_wins"][:10]
    plan["strategic_initiatives"] = plan["strategic_initiatives"][:10]

    return plan


def build_appendix(audit_results: dict[str, Any]) -> dict[str, Any]:
    """Build appendix section with full data."""
    appendix = {
        "methodology": get_methodology_text(),
        "full_issues": [],
        "technical_details": {},
    }

    # Collect all issues
    for audit_name, audit_data in audit_results.get("audits", {}).items():
        if audit_data:
            for issue in audit_data.get("issues", []):
                issue_copy = issue.copy()
                issue_copy["source"] = audit_name
                appendix["full_issues"].append(issue_copy)

    return appendix


def get_methodology_text() -> str:
    """Get methodology explanation text."""
    return """
This SEO Health Report was generated using automated analysis tools that evaluate:

1. **Technical Health (30%)**: Crawlability, page speed, mobile optimization,
   security, and structured data implementation.

2. **Content & Authority (35%)**: Content quality, E-E-A-T signals, topical
   coverage, and internal linking structure.

3. **AI Visibility (35%)**: How AI systems perceive, represent, and cite the
   brand, including knowledge graph presence and citation likelihood.

The composite score is calculated as a weighted average of these components,
with AI Visibility weighted equally to Content given its growing importance
in search discovery.

Scores are benchmarked against industry standards where available.
"""


def estimate_page_count(sections: list[dict[str, Any]]) -> int:
    """Estimate total page count based on sections."""
    # Rough estimates per section type
    estimates = {
        "cover": 1,
        "executive_summary": 1,
        "technical": 5,
        "content": 5,
        "ai_visibility": 5,
        "action_plan": 3,
        "appendix": 2,
    }

    total = 0
    for section in sections:
        section_type = section.get("type", "")
        total += estimates.get(section_type, 1)

    return total


def generate_docx(
    sections: list[dict[str, Any]],
    output_path: str,
    brand_colors: Optional[dict[str, str]] = None,
) -> tuple:
    """
    Generate DOCX document.

    Returns:
        Tuple of (success: bool, actual_path: str) - path may differ if fallback occurs
    """
    try:
        # Check if python-docx is available
        from docx import Document

        doc = Document()

        for section in sections:
            section_type = section.get("type", "")

            if section_type == "cover":
                content = section.get("content", {})
                # Add cover page
                doc.add_heading(content.get("title", "SEO Health Report"), 0)
                doc.add_paragraph(content.get("company_name", ""))
                doc.add_paragraph(content.get("date", ""))
                doc.add_paragraph(f"Prepared by: {content.get('prepared_by', '')}")
                doc.add_page_break()

            elif section_type == "executive_summary":
                content = section.get("content", {})
                doc.add_heading("Executive Summary", 1)
                doc.add_paragraph(content.get("headline", ""))

                # Score
                score_data = content.get("score_display", {})
                doc.add_paragraph(
                    f"Overall Score: {score_data.get('overall', 0)}/100 "
                    f"(Grade: {score_data.get('grade', 'N/A')})"
                )
                doc.add_paragraph(content.get("what_this_means", ""))
                doc.add_page_break()

            else:
                # Generic section
                doc.add_heading(section.get("title", "Section"), 1)
                content = section.get("content", {})

                if isinstance(content, dict):
                    for _key, value in content.items():
                        if isinstance(value, list):
                            for item in value[:5]:
                                if isinstance(item, dict):
                                    doc.add_paragraph(str(item), style="List Bullet")
                                else:
                                    doc.add_paragraph(str(item), style="List Bullet")
                doc.add_page_break()

        doc.save(output_path)
        return True, output_path

    except ImportError:
        # Fall back to markdown
        md_path = output_path.replace(".docx", ".md")
        success = generate_markdown(sections, md_path)
        return success, md_path if success else output_path
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        return False, output_path


def generate_pdf(
    sections: list[dict[str, Any]],
    output_path: str,
    brand_colors: Optional[dict[str, str]] = None,
) -> bool:
    """
    Generate PDF document using ReportLab.

    Args:
        sections: Report sections
        output_path: Output file path
        brand_colors: Optional brand colors

    Returns:
        True if successful
    """
    try:
        from xml.sax.saxutils import escape

        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

        from .pdf_components import (
            create_cover_page,
            create_findings_table,
            create_recommendations_list,
            create_section_header,
        )

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        story = []
        styles = getSampleStyleSheet()

        for section in sections:
            section_type = section.get("type", "")
            content = section.get("content", {})

            if section_type == "cover":
                story.extend(
                    create_cover_page(
                        title=escape(content.get("title", "SEO Health Report")),
                        company_name=escape(content.get("company_name", "")),
                        date=escape(content.get("date", "")),
                        logo_path=content.get("logo"),
                        brand_colors=brand_colors,
                    )
                )

            elif section_type == "executive_summary":
                story.extend(
                    create_section_header(
                        "Executive Summary", brand_colors=brand_colors
                    )
                )

                headline = escape(content.get("headline", ""))
                if headline:
                    story.append(Paragraph(f"<b>{headline}</b>", styles["Heading2"]))

                score_data = content.get("score_display", {})
                score_text = f"Overall Score: {score_data.get('overall', 0)}/100 (Grade: {escape(str(score_data.get('grade', 'N/A')))})"
                story.append(Paragraph(score_text, styles["Normal"]))
                story.append(Spacer(1, 0.2 * inch))

                what_means = escape(content.get("what_this_means", ""))
                if what_means:
                    story.append(Paragraph(what_means, styles["Normal"]))

                story.append(PageBreak())

            elif section_type in ["technical", "content", "ai_visibility"]:
                title = section.get("title", section_type.replace("_", " ").title())
                score = section.get("score")
                story.extend(create_section_header(escape(title), score, brand_colors))

                # Overview findings
                overview = content.get("overview", [])
                if overview:
                    story.append(create_findings_table(overview, brand_colors))
                    story.append(Spacer(1, 0.2 * inch))

                # Recommendations
                recs = content.get("recommendations", [])
                if recs:
                    story.append(
                        Paragraph("<b>Recommendations</b>", styles["Heading3"])
                    )
                    story.extend(create_recommendations_list(recs, brand_colors))

                story.append(PageBreak())

            elif section_type == "action_plan":
                story.extend(
                    create_section_header("Action Plan", brand_colors=brand_colors)
                )

                # Quick wins
                quick_wins = content.get("quick_wins", [])
                if quick_wins:
                    story.append(Paragraph("<b>Quick Wins</b>", styles["Heading3"]))
                    story.extend(create_recommendations_list(quick_wins, brand_colors))
                    story.append(Spacer(1, 0.2 * inch))

                # Prioritized tasks
                tasks = content.get("prioritized_tasks", [])
                if tasks:
                    story.append(
                        Paragraph("<b>Prioritized Tasks</b>", styles["Heading3"])
                    )
                    story.extend(create_recommendations_list(tasks[:10], brand_colors))

                story.append(PageBreak())

            elif section_type == "appendix":
                story.extend(
                    create_section_header("Appendix", brand_colors=brand_colors)
                )

                methodology = content.get("methodology", "")
                if methodology:
                    story.append(Paragraph("<b>Methodology</b>", styles["Heading3"]))
                    story.append(
                        Paragraph(
                            escape(methodology).replace("\n", "<br/>"), styles["Normal"]
                        )
                    )

        doc.build(story)
        return True

    except ImportError:
        logger.warning("ReportLab not installed. Falling back to markdown.")
        return generate_markdown(sections, output_path.replace(".pdf", ".md"))
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        return False


def generate_markdown(sections: list[dict[str, Any]], output_path: str) -> bool:
    """
    Generate Markdown document.

    Args:
        sections: Report sections
        output_path: Output file path

    Returns:
        True if successful
    """
    try:
        lines = []

        for section in sections:
            section_type = section.get("type", "")

            if section_type == "cover":
                content = section.get("content", {})
                lines.append(f"# {content.get('title', 'SEO Health Report')}")
                lines.append("")
                lines.append(f"**{content.get('company_name', '')}**")
                lines.append("")
                lines.append(f"*{content.get('date', '')}*")
                lines.append("")
                lines.append(f"Prepared by: {content.get('prepared_by', '')}")
                lines.append("")
                lines.append("---")
                lines.append("")

            elif section_type == "executive_summary":
                content = section.get("content", {})
                lines.append("## Executive Summary")
                lines.append("")
                lines.append(f"### {content.get('headline', '')}")
                lines.append("")

                score_data = content.get("score_display", {})
                lines.append(
                    f"**Overall Score: {score_data.get('overall', 0)}/100 "
                    f"(Grade: {score_data.get('grade', 'N/A')})**"
                )
                lines.append("")
                lines.append(content.get("what_this_means", ""))
                lines.append("")

                # Top actions
                lines.append("### Priority Actions")
                lines.append("")
                for action in content.get("top_actions", []):
                    lines.append(
                        f"- [{action.get('type', '')}] {action.get('action', '')}"
                    )
                lines.append("")
                lines.append("---")
                lines.append("")

            else:
                title = section.get("title", "Section")
                lines.append(f"## {title}")
                lines.append("")

                score = section.get("score")
                if score is not None:
                    lines.append(f"**Score: {score}/100**")
                    lines.append("")

                content = section.get("content", {})
                if isinstance(content, dict):
                    for key, value in content.items():
                        if isinstance(value, list) and value:
                            lines.append(f"### {key.replace('_', ' ').title()}")
                            lines.append("")
                            for item in value[:10]:
                                if isinstance(item, dict):
                                    item_str = (
                                        item.get("description", "")
                                        or item.get("action", "")
                                        or str(item)
                                    )
                                    lines.append(f"- {item_str}")
                                else:
                                    lines.append(f"- {item}")
                            lines.append("")

                lines.append("---")
                lines.append("")

        # Write file
        md_path = (
            output_path
            if output_path.endswith(".md")
            else output_path.replace(".docx", ".md").replace(".pdf", ".md")
        )

        with open(md_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

        return True

    except Exception as e:
        logger.error(f"Error generating Markdown: {e}")
        return False


__all__ = [
    "build_report_document",
    "build_technical_section",
    "build_content_section",
    "build_ai_visibility_section",
    "build_action_plan",
    "build_appendix",
    "generate_docx",
    "generate_pdf",
    "generate_markdown",
]
