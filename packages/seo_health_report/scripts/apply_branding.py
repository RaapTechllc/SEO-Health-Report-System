"""
Apply Branding

Apply company branding (logo, colors) to report documents.
"""

import os
import sys
from typing import Any, Optional

# Add parent directory to path for config import
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import get_config

# Get default colors from config
DEFAULT_COLORS = get_config().default_colors


def apply_branding(
    document_path: str,
    logo_file: Optional[str] = None,
    brand_colors: Optional[dict[str, str]] = None,
    company_name: Optional[str] = None,
) -> dict[str, Any]:
    """
    Apply branding to a document.

    Args:
        document_path: Path to the document to brand
        logo_file: Path to logo file (PNG/SVG)
        brand_colors: Custom brand colors
        company_name: Company name for text replacement

    Returns:
        Dict with branding result
    """
    result = {
        "success": False,
        "document_path": document_path,
        "logo_applied": False,
        "colors_applied": False,
        "error": None,
    }

    # Merge colors with defaults
    colors = DEFAULT_COLORS.copy()
    if brand_colors:
        colors.update(brand_colors)

    # Validate logo file
    if logo_file:
        if not os.path.exists(logo_file):
            result["error"] = f"Logo file not found: {logo_file}"
            return result

    # Determine document type
    if document_path.endswith(".docx"):
        result = apply_branding_docx(document_path, logo_file, colors, company_name)
    elif document_path.endswith(".pdf"):
        result["error"] = "PDF branding not yet implemented"
    elif document_path.endswith(".md"):
        # Markdown doesn't support branding in the same way
        result["success"] = True
        result["colors_applied"] = False
        result["logo_applied"] = False
    else:
        result["error"] = f"Unsupported document format: {document_path}"

    return result


def apply_branding_docx(
    document_path: str,
    logo_file: Optional[str],
    colors: dict[str, str],
    company_name: Optional[str],
) -> dict[str, Any]:
    """
    Apply branding to DOCX document.

    Args:
        document_path: Path to DOCX file
        logo_file: Path to logo file
        colors: Brand colors
        company_name: Company name

    Returns:
        Dict with branding result
    """
    result = {
        "success": False,
        "document_path": document_path,
        "logo_applied": False,
        "colors_applied": False,
        "error": None,
    }

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches

        doc = Document(document_path)

        # Apply logo to first paragraph/cover
        if logo_file and os.path.exists(logo_file):
            # Find the first paragraph and insert logo
            if doc.paragraphs:
                first_para = doc.paragraphs[0]
                run = first_para.insert_paragraph_before()
                run.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add logo
                try:
                    run = run.add_run()
                    run.add_picture(logo_file, width=Inches(2))
                    result["logo_applied"] = True
                except Exception as e:
                    result["error"] = f"Could not add logo: {e}"

        # Apply colors to headings
        # Note: Full color support requires modifying document styles
        # This is a simplified implementation

        primary_color = parse_hex_color(colors.get("primary", "#1a73e8"))

        for paragraph in doc.paragraphs:
            if paragraph.style and "Heading" in paragraph.style.name:
                for run in paragraph.runs:
                    run.font.color.rgb = primary_color

        result["colors_applied"] = True

        # Save document
        doc.save(document_path)
        result["success"] = True

    except ImportError:
        result["error"] = "python-docx package not installed"
    except Exception as e:
        result["error"] = str(e)

    return result


def parse_hex_color(hex_color: str):
    """
    Parse hex color string to RGBColor.

    Args:
        hex_color: Hex color string (e.g., "#1a73e8")

    Returns:
        RGBColor object or None
    """
    try:
        from docx.shared import RGBColor

        # Remove # if present
        hex_color = hex_color.lstrip("#")

        # Parse RGB values
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)

        return RGBColor(r, g, b)

    except Exception:
        return None


def generate_css_from_colors(colors: dict[str, str]) -> str:
    """
    Generate CSS variables from brand colors.

    Useful for HTML/web-based reports.

    Args:
        colors: Brand colors dict

    Returns:
        CSS string with variables
    """
    css_lines = [":root {"]

    for name, value in colors.items():
        css_name = name.replace("_", "-")
        css_lines.append(f"  --color-{css_name}: {value};")

    css_lines.append("}")

    return "\n".join(css_lines)


def get_score_color(score: float, colors: dict[str, str] = None) -> str:
    """
    Get appropriate color for a score value.

    Args:
        score: Score value (0-100)
        colors: Optional custom colors

    Returns:
        Hex color string
    """
    colors = colors or DEFAULT_COLORS
    config = get_config()

    if score >= config.grade_b_threshold:
        return colors.get("secondary", "#34a853")  # Green
    elif score >= config.grade_d_threshold:
        return colors.get("warning", "#fbbc04")  # Yellow
    else:
        return colors.get("danger", "#ea4335")  # Red


def create_color_scheme(primary: str, secondary: str = None) -> dict[str, str]:
    """
    Create a full color scheme from primary (and optionally secondary) colors.

    Args:
        primary: Primary brand color (hex)
        secondary: Optional secondary color (hex)

    Returns:
        Complete color scheme dict
    """
    scheme = DEFAULT_COLORS.copy()
    scheme["primary"] = primary

    if secondary:
        scheme["secondary"] = secondary

    return scheme


def validate_colors(colors: dict[str, str]) -> dict[str, Any]:
    """
    Validate color values.

    Args:
        colors: Colors to validate

    Returns:
        Dict with validation results
    """
    result = {"valid": True, "errors": []}

    import re

    hex_pattern = r"^#[0-9A-Fa-f]{6}$"

    for name, value in colors.items():
        if not re.match(hex_pattern, value):
            result["valid"] = False
            result["errors"].append(f"Invalid color for '{name}': {value}")

    return result


__all__ = [
    "DEFAULT_COLORS",
    "apply_branding",
    "apply_branding_docx",
    "parse_hex_color",
    "generate_css_from_colors",
    "get_score_color",
    "create_color_scheme",
    "validate_colors",
]
