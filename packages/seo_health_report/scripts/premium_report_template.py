"""
Premium SEO Health Report Template

Professional, branded report template focusing on AI visibility as the key differentiator.
Designed for high-value client deliverables.
"""

import os
from datetime import datetime
from typing import Optional


def generate_premium_docx_report(
    results: dict,
    overall: dict,
    technical: dict,
    content: dict,
    ai: dict,
    company_name: str,
    target_url: str,
    agency_name: str = "RaapTech",
    client_logo_path: Optional[str] = None,
    agency_logo_path: Optional[str] = None,
):
    """Generate a premium DOCX report with professional branding and AI focus."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()

        # === COVER PAGE ===
        # Header with agency logo (top right)
        if agency_logo_path and os.path.exists(agency_logo_path):
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header_para.add_run()
            run.add_picture(agency_logo_path, width=Inches(1.8))

        # Client logo (centered, prominent)
        if client_logo_path and os.path.exists(client_logo_path):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.space_before = Pt(72)  # 1 inch from top
            run = logo_para.add_run()
            run.add_picture(client_logo_path, width=Inches(4))

        # Main title
        title = doc.add_heading("SEO HEALTH REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(36)
        title_run.font.color.rgb = RGBColor(26, 54, 93)  # Professional blue

        # Subtitle with company name
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.space_before = Pt(24)
        run = subtitle.add_run(f"{company_name}")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(43, 108, 176)

        # Score highlight box
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_para.space_before = Pt(36)
        score_run = score_para.add_run(
            f"OVERALL SCORE: {overall['overall_score']}/100 (Grade: {overall['grade']})"
        )
        score_run.bold = True
        score_run.font.size = Pt(18)

        # Grade color coding
        grade_colors = {
            "A": RGBColor(34, 139, 34),  # Green
            "B": RGBColor(50, 205, 50),  # Light green
            "C": RGBColor(255, 165, 0),  # Orange
            "D": RGBColor(255, 69, 0),  # Red orange
            "F": RGBColor(220, 20, 60),  # Crimson
        }
        score_run.font.color.rgb = grade_colors.get(overall["grade"], RGBColor(0, 0, 0))

        # AI Visibility callout (key differentiator)
        ai_callout = doc.add_paragraph()
        ai_callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ai_callout.space_before = Pt(24)
        ai_run = ai_callout.add_run("✨ INCLUDING EXCLUSIVE AI VISIBILITY ANALYSIS ✨")
        ai_run.bold = True
        ai_run.font.size = Pt(14)
        ai_run.font.color.rgb = RGBColor(138, 43, 226)  # Purple for AI

        # Date and agency info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.space_before = Pt(48)
        info_para.add_run(f"Prepared by: {agency_name}\n").font.size = Pt(12)
        info_para.add_run(
            f"Date: {datetime.now().strftime('%B %d, %Y')}\n"
        ).font.size = Pt(12)
        info_para.add_run(f"Website: {target_url}").font.size = Pt(10)

        doc.add_page_break()

        # === EXECUTIVE SUMMARY ===
        doc.add_heading("Executive Summary", level=1)

        # Grade description
        grade_descriptions = {
            "A": "🏆 EXCELLENT - Your SEO health is outstanding! You're ahead of 90% of competitors.",
            "B": "✅ GOOD - Strong SEO foundation with strategic opportunities for growth.",
            "C": "⚠️ NEEDS ATTENTION - Several critical areas require immediate focus.",
            "D": "🚨 POOR - Significant SEO issues are limiting your online visibility.",
            "F": "🔴 CRITICAL - Urgent action required to prevent further ranking decline.",
        }

        summary_para = doc.add_paragraph()
        summary_para.add_run(
            f"{grade_descriptions.get(overall['grade'], 'Assessment complete')}\n\n"
        )

        # Key insight about AI visibility
        ai_score = ai.get("score", 0)
        if ai_score >= 80:
            ai_insight = "🎯 COMPETITIVE ADVANTAGE: Your brand has strong AI visibility, positioning you ahead of competitors who aren't optimizing for AI search systems."
        elif ai_score >= 60:
            ai_insight = "📈 OPPORTUNITY: Moderate AI visibility with significant room for improvement. This is your chance to get ahead of competitors."
        else:
            ai_insight = "🚀 UNTAPPED POTENTIAL: Low AI visibility represents a major growth opportunity. Most competitors aren't optimizing for AI - you can lead the market."

        ai_para = doc.add_paragraph()
        ai_run = ai_para.add_run(ai_insight)
        ai_run.font.color.rgb = RGBColor(138, 43, 226)
        ai_run.bold = True

        # Score breakdown table
        doc.add_heading("Performance Breakdown", level=2)
        table = doc.add_table(rows=5, cols=4)
        table.style = "Table Grid"

        # Headers
        headers = table.rows[0].cells
        headers[0].text = "SEO Component"
        headers[1].text = "Your Score"
        headers[2].text = "Industry Avg"
        headers[3].text = "Impact Weight"

        # Make headers bold
        for cell in headers:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        rows_data = [
            ("🔧 Technical SEO", f"{technical.get('score', 0)}/100", "65/100", "30%"),
            (
                "📝 Content & Authority",
                f"{content.get('score', 0)}/100",
                "58/100",
                "35%",
            ),
            ("🤖 AI Visibility", f"{ai.get('score', 0)}/100", "25/100", "35%"),
            ("📊 OVERALL SCORE", f"{overall['overall_score']}/100", "49/100", "100%"),
        ]

        for i, (component, score, industry, weight) in enumerate(rows_data, 1):
            row = table.rows[i].cells
            row[0].text = component
            row[1].text = score
            row[2].text = industry
            row[3].text = weight

            # Highlight overall score row
            if i == 4:
                for cell in row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        doc.add_page_break()

        # === AI VISIBILITY SECTION (FEATURED) ===
        ai_heading = doc.add_heading("🤖 AI Visibility Analysis", level=1)
        ai_heading.runs[0].font.color.rgb = RGBColor(138, 43, 226)

        # AI intro paragraph
        ai_intro = doc.add_paragraph()
        ai_intro.add_run(
            "This exclusive analysis evaluates how your brand appears in AI-powered search systems - a critical factor that 95% of businesses ignore.\n\n"
        ).bold = True

        ai_systems_para = doc.add_paragraph()
        ai_systems_para.add_run(
            "We tested your brand visibility across the top AI systems:\n"
        )
        ai_systems_para.add_run("• Google AI Search (Bard/Gemini)\n")
        ai_systems_para.add_run("• Perplexity AI Search\n")
        ai_systems_para.add_run("• OpenAI Search (ChatGPT)\n\n")

        # AI score with context
        ai_score_para = doc.add_paragraph()
        ai_score_para.add_run(
            f"Your AI Visibility Score: {ai.get('score', 0)}/100\n"
        ).bold = True

        if ai.get("score", 0) >= 70:
            ai_context = "🎯 EXCELLENT: Your brand consistently appears in AI responses with accurate information."
        elif ai.get("score", 0) >= 50:
            ai_context = "📈 GOOD: Your brand appears in some AI responses but needs optimization for consistency."
        else:
            ai_context = "🚀 OPPORTUNITY: Significant potential to improve AI visibility and get ahead of competitors."

        ai_score_para.add_run(ai_context)

        # AI components breakdown
        if ai.get("components"):
            doc.add_heading("AI Performance Components", level=2)

            ai_components = [
                ("ai_presence", "Brand Mention Rate", 25),
                ("accuracy", "Information Accuracy", 20),
                ("parseability", "AI Crawlability", 15),
                ("knowledge_graph", "Knowledge Graph Presence", 15),
                ("citation_likelihood", "Citation Potential", 15),
                ("sentiment", "Brand Sentiment", 10),
            ]

            for comp_key, comp_name, max_points in ai_components:
                comp_data = ai.get("components", {}).get(comp_key, {})
                if comp_data:
                    score = comp_data.get("score", 0)
                    doc.add_heading(f"{comp_name}: {score}/{max_points}", level=3)

                    # Add findings if available
                    findings = comp_data.get("findings", [])
                    if findings:
                        for finding in findings[:3]:  # Top 3 findings
                            doc.add_paragraph(f"• {finding}", style="List Bullet")

        doc.add_page_break()

        # === TECHNICAL SEO SECTION ===
        doc.add_heading("🔧 Technical SEO Analysis", level=1)
        doc.add_paragraph(
            f"Technical Foundation Score: {technical.get('score', 0)}/100"
        )

        # Technical components
        if technical.get("components"):
            tech_components = [
                ("crawlability", "Site Crawlability", 20),
                ("speed", "Page Speed & Core Web Vitals", 25),
                ("security", "Security & HTTPS", 10),
                ("mobile", "Mobile Optimization", 15),
                ("structured_data", "Structured Data", 15),
            ]

            for comp_key, comp_name, max_points in tech_components:
                comp_data = technical.get("components", {}).get(comp_key, {})
                if comp_data:
                    score = comp_data.get("score", 0)
                    doc.add_heading(f"{comp_name}: {score}/{max_points}", level=2)

        doc.add_page_break()

        # === CONTENT & AUTHORITY SECTION ===
        doc.add_heading("📝 Content & Authority Analysis", level=1)
        doc.add_paragraph(f"Content Authority Score: {content.get('score', 0)}/100")

        # Content components
        if content.get("components"):
            content_components = [
                ("content_quality", "Content Quality", 25),
                ("eeat", "E-E-A-T Signals", 20),
                ("topical_authority", "Topical Authority", 15),
                ("backlinks", "Backlink Profile", 15),
                ("internal_links", "Internal Link Structure", 10),
            ]

            for comp_key, comp_name, max_points in content_components:
                comp_data = content.get("components", {}).get(comp_key, {})
                if comp_data:
                    score = comp_data.get("score", 0)
                    doc.add_heading(f"{comp_name}: {score}/{max_points}", level=2)

        doc.add_page_break()

        # === ACTION PLAN ===
        doc.add_heading("🎯 Strategic Action Plan", level=1)

        # Priority recommendations
        doc.add_heading("High Priority Actions (Next 30 Days)", level=2)

        high_priority = []
        if ai.get("score", 0) < 60:
            high_priority.append(
                "🤖 Optimize content for AI systems - create comprehensive, factual pages about your services"
            )
        if technical.get("score", 0) < 70:
            high_priority.append(
                "🔧 Fix critical technical issues - improve page speed and mobile optimization"
            )
        if content.get("score", 0) < 60:
            high_priority.append(
                "📝 Enhance content authority - add author bios and expertise signals"
            )

        for action in high_priority[:5]:  # Top 5 actions
            doc.add_paragraph(action, style="List Bullet")

        # Medium priority
        doc.add_heading("Medium Priority Actions (Next 90 Days)", level=2)
        medium_actions = [
            "📊 Implement structured data markup for better AI understanding",
            "🔗 Build high-quality backlinks from industry authorities",
            "📱 Optimize for mobile-first indexing",
            "🎯 Create topic cluster content strategy",
        ]

        for action in medium_actions:
            doc.add_paragraph(action, style="List Bullet")

        # Footer with branding
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"© {datetime.now().year} {agency_name} | Premium SEO Health Report for {company_name}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)

        # Save document
        project_root = os.path.dirname(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        )
        reports_dir = os.path.join(project_root, "reports")
        os.makedirs(reports_dir, exist_ok=True)

        output_path = os.path.join(
            reports_dir,
            f"{company_name.replace(' ', '-')}-Premium-SEO-Report-{datetime.now().strftime('%Y-%m-%d')}.docx",
        )
        doc.save(output_path)

        print(f"\n[PREMIUM REPORT GENERATED] {output_path}")
        print("Professional branded report with AI visibility focus")
        print("Ready for high-value client delivery")

        return output_path

    except ImportError:
        print(
            "\n[ERROR] python-docx not installed. Run 'pip install python-docx' for premium report generation."
        )
        return None
    except Exception as e:
        print(f"\n[ERROR] Error generating premium report: {e}")
        import traceback

        traceback.print_exc()
        return None
